from docx import Document

# Criação do documento Word
doc = Document()
doc.add_heading('Tabela de Terapias', level=1)

# Adicionar tabela com cabeçalhos
table = doc.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Data'
hdr_cells[1].text = 'Hora'
hdr_cells[2].text = 'Terapia'

# Dados da tabela completos
data_rows = [
    ['04/07/24 - qui', '11:00 - 12:00', 'Exercício de marcha com auxílio'],
    ['05/07/24 - sex', '10:30 - 11:30', 'Fisioterapia respiratória - Expansão torácica'],
    ['06/07/24 - sáb', '10:30 - 11:30', 'Exercício de marcha com barras paralelas'],
    ['07/07/24 - dom', '12:00 - 13:00', 'Exercícios de controle respiratório'],
    ['08/07/24 - seg', '17:00 - 18:00', 'Treino de marcha em superfícies irregulares'],
    ['09/07/24 - ter', '10:30 - 11:30', 'Exercícios respiratórios - PEP e Flutter'],
    ['10/07/24 - qua', 'NA', ''],
    ['11/07/24 - qui', '11:00 - 12:00', 'Marcha com transferência de peso'],
    ['12/07/24 - sex', '10:30 - 11:30', 'Exercício de respiração diafragmática'],
    ['13/07/24 - sáb', '10:30 - 11:30', 'Treino de marcha com obstáculos'],
    ['14/07/24 - dom', '12:00 - 13:00', 'Expansão pulmonar com incentivador'],
    ['15/07/24 - seg', '17:00 - 18:00', 'Exercícios de equilíbrio dinâmico'],
    ['16/07/24 - ter', '10:30 - 11:30', 'Fisioterapia respiratória - Mobilização torácica'],
    ['17/07/24 - qua', 'NA', ''],
    ['18/07/24 - qui', '11:00 - 12:00', 'Marcha com suporte parcial de peso'],
    ['19/07/24 - sex', '10:30 - 11:30', 'Exercícios respiratórios - Inspirometria'],
    ['20/07/24 - sáb', '10:30 - 11:30', 'Treino de subida e descida de escadas'],
    ['21/07/24 - dom', '12:00 - 13:00', 'Expansão pulmonar ativa assistida'],
    ['22/07/24 - seg', '17:00 - 18:00', 'Marcha com estimulação funcional elétrica'],
    ['23/07/24 - ter', '10:30 - 11:30', 'Fisioterapia respiratória - Desobstrução brônquica'],
    ['24/07/24 - qua', 'NA', ''],
    ['25/07/24 - qui', '11:00 - 12:00', 'Exercícios de transferência de peso na marcha'],
    ['26/07/24 - sex', '10:30 - 11:30', 'Expansão torácica com padrões respiratórios'],
    ['27/07/24 - sáb', '10:30 - 11:30', 'Marcha em circuito com variação de velocidade'],
    ['28/07/24 - dom', '12:00 - 13:00', 'Exercícios respiratórios ativos'],
    ['29/07/24 - seg', '17:00 - 18:00', 'Treino de marcha com desvio de direção'],
    ['30/07/24 - ter', '10:30 - 11:30', 'Fisioterapia respiratória com técnicas de higiene brônquica'],
    ['31/07/24 - qua', 'NA', ''],
    ['01/08/24 - qui', '11:00 - 12:00', 'Marcha com uso de andador'],
    ['02/08/24 - sex', '10:30 - 11:30', 'Exercícios de respiração com uso de bola de reabilitação'],
    ['03/08/24 - sáb', '10:30 - 11:30', 'Marcha com obstáculos dinâmicos'],
]

# Adicionando dados à tabela
for row in data_rows:
    cells = table.add_row().cells
    cells[0].text = row[0]
    cells[1].text = row[1]
    cells[2].text = row[2]

# Salvar o arquivo
doc.save('Tabela_de_Terapias.docx')